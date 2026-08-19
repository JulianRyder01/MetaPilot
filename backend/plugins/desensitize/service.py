"""脱敏插件核心服务：本地模型（ollama）识别敏感信息 + 开放替换/涂黑工具集。

流程（与前端约定，共三阶段）：
1. analyze：把文本/文档交给本地模型（ollama，默认 qwen3.5:4b 可改）输出 JSON 敏感清单
   [{"value": 原文敏感串, "type": 敏感类型}]，工具对这些 value 在原文中精确定位并返回
   start/end（供前端高亮「标出所有敏感信息」）。
2. apply：用户确认要脱敏的条目后，工具按确认条目对文本/markdown 用 ■（黑块）替换原文，
   并返回替换后的文本与发生替换的区间（spans）。
3. 文件（PDF/图片）：工具用黑色块涂黑敏感内容所在区域，输出新文件供下载。

「开放工具集」说明：文本替换（replace:text / replace:text:label）、PDF 涂黑（redact:pdf）、
图片涂黑（redact:image）均为插件内置的脱敏工具；AI 只负责输出 JSON 识别结果，具体替换由
工具执行。识别出的敏感类型由模型自由返回，本模块仅提供「类型别名 → 展示名」的可扩展映射，
未命中的类型原样展示（不写死、可动态扩充）。
"""
from __future__ import annotations

import asyncio
import io
import json
import re
import threading
import uuid
from typing import Optional

from app.services.ai_config import AIConfig
from app.services.ollama import OllamaClient, OllamaError

# 敏感类型别名 → 展示名（可动态扩充：未知类型原样返回，不写死）
_TYPE_LABELS: dict[str, str] = {
    "phone": "手机号", "mobile": "手机号", "tel": "电话",
    "id": "证件号", "id_card": "身份证", "passport": "护照",
    "email": "邮箱", "name": "姓名", "address": "地址", "bank": "银行卡",
    "card": "银行卡", "bank_card": "银行卡", "account": "账号", "plate": "车牌", "ip": "IP",
    "credit_code": "统一社会信用代码", "tax": "税号", "date": "出生日期",
    "other": "其他",
}

# 规则正则引擎：对格式明确的编码型敏感信息（证件号/手机/座机/银行卡/邮箱/IP/车牌）
# 作可靠识别与定位，不依赖本地模型；模型对长文本识别弱/输出空时由规则兜底。
# 顺序即优先级：身份证先占位，银行卡跳过与其重叠的 18 位区间（防止身份证被当银行卡）。
_RE_PATTERNS: dict[str, str] = {
    "id_card": r"(?<!\d)\d{17}[\dXx](?!\d)",            # 身份证 18 位
    "phone": r"(?<!\d)1[3-9]\d{9}(?!\d)",                # 手机号
    "tel": r"(?<!\d)0\d{2,3}-?\d{7,8}(?!\d)",            # 座机
    "email": r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+",           # 邮箱
    "ip": r"(?<!\d)(?:\d{1,3}\.){3}\d{1,3}(?!\d)",      # IP
    "plate": r"[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-HJ-NP-Z][A-HJ-NP-Z0-9]{5,6}",  # 车牌
    "bank_card": r"(?<!\d)\d{16,19}(?!\d)",              # 银行卡（16-19 位，排除身份证）
    "passport": r"(?<!\d)[A-Za-z]\d{8}(?!\d)",            # 护照（字母+8 位数字）
    "credit_code": r"[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}",  # 统一社会信用代码（18 位，排除身份证）
}

# 超过该字符数的长文本不再调用本地模型识别（已实测：qwen3.5:4b 这类小模型对长文本
# 输出不可靠——发散生成到上下文上限或被截断为空，耗时久且无有效结果），仅用规则引擎
# 可靠识别编码型 + 合同语境姓名类敏感信息，保证秒回与可用。
_AI_MAX_CHARS = 800

# 姓名（合同/签字语境，如“甲方（委托方）：林浩然”），限定在法定/委托/签字语境，避免误抓地址等。
_NAME_PATTERN = re.compile(
    r"(?:甲方|乙方)[（(]?(?:委托方|服务方|签字|法定代表人)[）)]?\s*[：:]\s*([\u4e00-\u9fa5]{2,4})"
)

# 百家姓（常见单姓，通用姓名识别的首字），拼进正则字符类
_SURNAMES = "赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜戚谢邹喻柏水窦章云苏潘葛奚范彭郎鲁韦昌马苗凤花方俞任袁柳酆鲍史唐费廉岑薛雷贺倪汤滕殷罗毕郝邬安常乐于时傅皮卞齐康伍余元卜顾孟平黄和穆萧尹姚邵湛汪祁毛禹狄米贝明臧计伏成戴谈宋茅庞熊纪舒屈项祝董梁杜阮蓝闵席季麻强贾路娄危江童颜郭梅盛林刁钟徐邱骆高夏蔡田樊胡凌霍虞万支柯昝管卢莫经房裘缪干解应宗丁宣贲邓郁单杭洪包诸左石崔吉钮龚程嵇邢滑裴陆荣翁荀羊於惠甄曲家封芮羿储靳汲邴糜松井段富巫乌焦巴弓牧隗山谷车侯宓蓬全郗班仰秋仲伊宫宁仇栾暴甘斜厉戎祖武符刘景詹束龙叶幸司韶郜黎蓟溥印宿白怀蒲邰从鄂索咸籍赖卓蔺屠蒙池乔阴郁胥能苍双闻莘党翟谭贡劳逄姬申扶堵冉宰郦雍郤璩桑桂濮牛寿通边扈燕冀郏浦尚农温别庄晏柴瞿阎充慕连茹习宦艾鱼容向古易慎戈廖庾终暨居衡步都耿满弘匡国文寇广禄阙东欧殳沃利蔚越夔隆师巩厍聂晁勾敖融冷訾辛阚那简饶空曾毋沙乜养鞠须丰巢关蒯相查后荆红游竺权逯盖益桓公"

# 通用姓名：百家姓开头 + 1-2 个汉字，前后非汉字/字母/数字（独立词形，减少误抓）
_GEN_NAME_PATTERN = re.compile(
    r"(?<![\u4e00-\u9fa5A-Za-z0-9])([" + _SURNAMES + r"])([\u4e00-\u9fa5]{1,2})(?![\u4e00-\u9fa5])"
)

# 地址：省市区 / 路街大道 / 号栋室等（长度过滤在调用侧，避免短误伤）
_ADDR_PATTERN = re.compile(
    r"(?<![\u4e00-\u9fa5A-Za-z0-9])"
    r"([\u4e00-\u9fa5]{2,8}?(?:省|自治区|市|自治州|地区))?\s*"
    r"([\u4e00-\u9fa5]{2,10}?(?:市|区|县|镇))?\s*"
    r"[\u4e00-\u9fa5]{1,16}?(?:路|街|大道|巷|弄)\s*[\u4e00-\u9fa5A-Za-z0-9]{0,5}?(?:号|栋|幢|室|单元)?"
)

# 黑块字符（替换/涂黑的统一视觉单元）
BLACK = "█"


def type_label(t: str) -> str:
    return _TYPE_LABELS.get(t or "", t or "其他")


def _parse_json(text: str) -> dict:
    """稳健解析模型 JSON 输出：剥离 ```json 围栏、截取首个 { ... } 平衡块。"""
    txt = text.strip()
    if txt.startswith("```"):
        txt = re.sub(r"^```[a-zA-Z]*\s*", "", txt)
        txt = re.sub(r"\s*```$", "", txt).strip()
    try:
        return json.loads(txt)
    except Exception:
        pass
    # 截取平衡花括号
    start = txt.find("{")
    if start >= 0:
        depth = 0
        for i in range(start, len(txt)):
            if txt[i] == "{":
                depth += 1
            elif txt[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(txt[start:i + 1])
                    except Exception:
                        break
    return {}


# ---------- 本机 Tesseract OCR（扫描件 PDF / 图片识别） ----------

# tesseract 可不在 PATH，取常见安装位置兜底；未安装时 OCR 能力降级为不可用（不崩溃）。
_TESSERACT_CANDIDATES = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Tesseract-OCR\tesseract.exe",
    "/usr/bin/tesseract",
    "/usr/local/bin/tesseract",
)
_OCR_LANG = "chi_sim+eng"  # 简体中文 + 英文


def _tesseract_cmd() -> str:
    """探测 tesseract 可执行路径（PATH 优先，其次常见安装位置）。"""
    import os
    import shutil

    exe = shutil.which("tesseract") or ""
    if not exe:
        for c in _TESSERACT_CANDIDATES:
            if os.path.exists(c):
                exe = c
                break
    return exe


def _ensure_tesseract() -> bool:
    """定位并配置 pytesseract 的 tesseract_cmd；返回是否可用。"""
    cmd = _tesseract_cmd()
    if not cmd:
        return False
    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = cmd
        return True
    except Exception:
        return False


def _doc_binary_text(data: bytes) -> str:
    """尽力从旧 .doc（OLE 复合文档二进制）提取可读文本；无专业解析库时的降级探测。

    能抠出 UTF-16LE/GBK/Latin 下的连续可打印段；提取不到清晰文本时由调用方报错提示。
    """
    best = ""
    for enc in ("utf-16-le", "utf-8", "gb18030", "latin-1"):
        try:
            s = data.decode(enc, errors="ignore")
        except Exception:
            continue
        chunks = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9，。：；、（）()：；\-\s]{4,}", s)
        text = "\n".join(chunks).strip()
        if len(text) > len(best):
            best = text
    return best


# 批量导入任务注册表（taskId → task）。task.files 的 stage（queued→extracting→analyzing→done/error）
# 用于前端展示“文件内阶段级”进度；单进程 FastAPI 内共享即可。
_IMPORT_TASKS: dict[str, dict] = {}
_IMPORT_LOCK = threading.Lock()


class DesensitizeService:
    """脱敏服务：识别 + 替换/涂黑工具端点所需，经 request.app.state.desensitize 取用。"""

    def __init__(self, ollama: Optional[OllamaClient] = None, config: Optional[AIConfig] = None):
        self.ollama = ollama or OllamaClient(config=config)
        self.config = config or self.ollama.config if hasattr(self.ollama, "config") else None

    # ---------------- 批量导入（进度阶段式） ----------------

    def start_import(self, files: list[tuple[str, bytes]], folder_id: str, model: str = "") -> dict:
        """批量导入：接收 [(文件名, bytes)]，后台线程逐文件提取+识别，返回 taskId。

        后台线程处理期间前端经 get_import/{taskId} 轮询，每文件 stage 为
        queued → extracting → analyzing → done / error（文件内阶段级进度）。
        """
        tid = uuid.uuid4().hex[:12]
        task = {
            "id": tid, "status": "running", "total": len(files), "done": 0,
            "folderId": folder_id,
            "files": [{"name": n, "stage": "queued", "progress": 0.0} for n, _ in files],
        }
        with _IMPORT_LOCK:
            _IMPORT_TASKS[tid] = task
        threading.Thread(target=self._run_import, args=(task, files, model), daemon=True).start()
        return {"taskId": tid, "total": len(files),
                "files": [{"name": n, "stage": "queued", "progress": 0.0} for n, _ in files]}

    def _run_import(self, task: dict, files: list, model: str) -> None:
        """后台线程：逐文件提取→识别，更新 task.files 状态；单文件失败不中断其余。"""
        for i, (name, data) in enumerate(files):
            f = task["files"][i]
            try:
                f["stage"] = "extracting"; f["progress"] = 0.25
                ex = self._extract_file(name, data)
                f.update({"kind": ex.get("kind"), "scanned": ex.get("scanned", False),
                          "text": ex.get("text", "")})
                if not ex.get("text", "").strip():
                    f.update({"stage": "done", "progress": 1.0, "count": 0, "items": []})
                    task["done"] += 1
                    continue
                f["stage"] = "analyzing"; f["progress"] = 0.6
                items = asyncio.run(self.analyze_text(ex["text"], model))
                f.update({"items": items["items"], "count": items.get("count", 0),
                          "model": items.get("model", model), "stage": "done", "progress": 1.0})
            except Exception as e:
                f.update({"stage": "error", "error": str(e)})
            task["done"] += 1
        task["status"] = "done"

    def _extract_file(self, name: str, data: bytes) -> dict:
        """按扩展名分派文本提取（PDF/图片/办公文档/纯文本）。"""
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext == "pdf":
            return self.extract_pdf_text(data)
        if ext in ("png", "jpg", "jpeg", "bmp", "webp", "gif"):
            return self.extract_image_text(data)
        return self.extract_doc_text(data, ext)

    def get_import(self, task_id: str) -> dict:
        """查询批量导入任务状态（每文件阶段/进度/识别结果；text 截断防超大响应）。"""
        with _IMPORT_LOCK:
            task = _IMPORT_TASKS.get(task_id)
        if not task:
            raise KeyError(f"任务不存在: {task_id}")
        files = []
        for f in task["files"]:
            item = {k: v for k, v in f.items()}
            if isinstance(item.get("text"), str):
                item["text"] = item["text"][:3000]
            files.append(item)
        return {"id": task["id"], "status": task["status"], "total": task["total"],
                "done": task["done"], "folderId": task["folderId"], "files": files}

    # ---------------- 工具：文本替换（replace:text） ----------------

    @staticmethod
    def mask_text(text: str, items: list[dict]) -> tuple[str, list[dict]]:
        """把确认条目的敏感串替换为黑色块；返回 (替换后文本, 实际替换区间 spans)。

        全局替换：对每条确认条目按 value 在原文中查找【所有】出现处（不限于该条目的
        start/end），同名敏感串在文中出现多少次就替换多少次；重叠区间合并，保证不重复替换。
        """
        if not items or not text:
            return text, []
        spans: list[tuple[int, int, dict]] = []
        for it in items:
            v = str(it.get("value") or "").strip()
            if not v:
                continue
            # 全局替换：同一条目 value 在全文所有出现处都涂黑（不受单条 start/end 限制）
            pos = 0
            while True:
                idx = text.find(v, pos)
                if idx < 0:
                    break
                spans.append((idx, idx + len(v), it))
                pos = idx + len(v)
        if not spans:
            return text, []
        spans.sort(key=lambda x: (x[0], -x[1]))
        merged: list[tuple[int, int, dict]] = []
        for s, e, it in spans:
            if merged and s < merged[-1][1]:
                continue
            merged.append((s, e, it))
        out: list[str] = []
        last = 0
        result_spans: list[dict] = []
        for s, e, it in merged:
            out.append(text[last:s])
            value = text[s:e] or str(it.get("value"))
            out.append(BLACK * (e - s))
            result_spans.append({"start": s, "end": e, "value": value,
                                 "type": it.get("type", ""), "typeLabel": type_label(it.get("type", ""))})
            last = e
        out.append(text[last:])
        return "".join(out), result_spans

    # ---------------- 识别 ----------------

    def _system_prompt(self) -> str:
        return (
            "你是本地数据脱敏助手。阅读用户提供的文本，找出所有敏感信息，例如：姓名、手机号、"
            "身份证号、护照、邮箱、家庭/公司地址、银行卡/银行账号、车牌号、IP 地址等个人信息。\n"
            "只输出 JSON，不要任何其它文字或解释。格式："
            '{"items":[{"value":"敏感串，必须与原文完全一致","type":"类型英文标识"}]}。\n'
            "要求：value 必须逐字符与原文一致（用于精确定位并替换），不要改写、截断或省略；"
            "同一种若出现多次可各列一条；没有敏感信息时输出 {\"items\":[]}。"
        )

    async def analyze_text(self, text: str, model: str = "") -> dict:
        """识别敏感信息：先跑规则正则引擎，再用本地模型（AI）补充语义型，合并去重。

        规则引擎对证件号/手机号/银行卡/邮箱/IP/座机/车牌等编码型敏感信息可靠识别定位；
        AI 补充姓名/地址等语义型。AI 不可用、失败或对长文本输出为空时，规则结果照常返回
        （不再因模型失败/空输出而整体识别为 0 条）。返回条目含 start/end/found（供高亮）。
        """
        if not text or not text.strip():
            return {"items": [], "count": 0}
        model = model or self.ollama.llm_model
        # 1) 规则引擎（毫秒级、可靠）
        reg_items = self._regex_items(text)
        # 2) AI 补充（失败/空不阻断，规则兜底；长文本跳过——4B 模型对长文本输出不可靠且慢）
        ai_items: list[dict] = []
        ai_model = model
        if len(text) <= _AI_MAX_CHARS:
            try:
                messages = [
                    {"role": "system", "content": self._system_prompt()},
                    {"role": "user", "content": text},
                ]
                res = await self.ollama.chat(messages, model=model, json_mode=True, temperature=0)
                ai_model = res.get("model", model)
                data = _parse_json(res.get("content", ""))
                raw_items = data.get("items") if isinstance(data.get("items"), list) else []
                for it in raw_items:
                    if not isinstance(it, dict):
                        continue
                    v = str(it.get("value") or "").strip()
                    if not v:
                        continue
                    t = str(it.get("type") or "").strip()
                    st = it.get("start")
                    en = it.get("end")
                    found, s, e = self._locate(text, v, st, en)
                    ai_items.append({"value": v, "type": t, "typeLabel": type_label(t),
                                     "start": s, "end": e, "found": found, "source": "ai"})
            except (OllamaError, RuntimeError):
                ai_items = []  # AI 不可用 → 仅用规则结果，不整体失败
        # 3) 合并规则与 AI（按区间/值去重，排序）
        items = self._merge_items(reg_items, ai_items, text)
        return {"items": items, "count": len(items), "model": ai_model,
                "regex": len(reg_items), "ai": len(ai_items)}

    @staticmethod
    def _regex_items(text: str) -> list[dict]:
        """规则正则引擎：对编码型(证件/手机/银行/邮箱/IP/车牌/护照/信用代码)、
        姓名(合同语境+百家姓通用)、地址做可靠定位，返回带 start/end 的条目。

        身份证最优先占位；信用代码/银行卡跳过与其重叠的 18 位身份证，避免误判。
        """
        items: list[dict] = []
        covered: list[tuple[int, int]] = []
        ordered = ["id_card", "phone", "tel", "email", "ip", "plate", "bank_card",
                   "passport", "credit_code"]
        for typ in ordered:
            pat = _RE_PATTERNS.get(typ)
            if not pat:
                continue
            for m in re.finditer(pat, text):
                s, e = m.span()
                if any(s < c[1] and e > c[0] for c in covered):
                    continue
                items.append({"value": text[s:e], "type": typ, "typeLabel": type_label(typ),
                              "start": s, "end": e, "found": True, "source": "regex"})
                covered.append((s, e))
        # 姓名：合同/签字语境 + 通用百家姓；同一姓名去重保留首处
        seen_names: set[str] = set()
        for m in _NAME_PATTERN.finditer(text):
            name = m.group(1)
            s, e = m.span(1)
            if name in seen_names or any(s < c[1] and e > c[0] for c in covered):
                continue
            seen_names.add(name)
            items.append({"value": name, "type": "name", "typeLabel": type_label("name"),
                          "start": s, "end": e, "found": True, "source": "regex"})
            covered.append((s, e))
        for m in _GEN_NAME_PATTERN.finditer(text):
            name = m.group(0)
            s, e = m.span()
            if name in seen_names or any(s < c[1] and e > c[0] for c in covered):
                continue
            seen_names.add(name)
            items.append({"value": name, "type": "name", "typeLabel": type_label("name"),
                          "start": s, "end": e, "found": True, "source": "regex"})
            covered.append((s, e))
        # 地址（省市区/路街/号室等）：截取整段，长度≥6 防过度误抓
        for m in _ADDR_PATTERN.finditer(text):
            s, e = m.span()
            val = text[s:e]
            if len(val) < 6 or any(s < c[1] and e > c[0] for c in covered):
                continue
            items.append({"value": val, "type": "address", "typeLabel": type_label("address"),
                          "start": s, "end": e, "found": True, "source": "regex"})
            covered.append((s, e))
        items.sort(key=lambda x: x["start"])
        return items

    @staticmethod
    def _merge_items(reg: list[dict], ai: list[dict], text: str) -> list[dict]:
        """合并规则项与 AI 项：与已占用区间重叠或 value 相同的 AI 项跳过，其余并入后排序。"""
        merged = list(reg)
        spans = [(it["start"], it["end"]) for it in reg]
        seen_values = {it.get("value") for it in reg}
        for it in ai:
            v = it.get("value")
            if not v or v in seen_values:
                continue
            s, e = it.get("start", 0), it.get("end", 0)
            if any(s < c[1] and e > c[0] for c in spans):
                continue
            merged.append(it)
            spans.append((s, e))
            seen_values.add(v)
        merged.sort(key=lambda x: x["start"])
        return merged

    @staticmethod
    def _locate(text: str, value: str, start=None, end=None) -> tuple[bool, int, int]:
        """返回 (是否定位到, 首处 start, end)；优先给定区间校验，否则子串查找，再宽松匹配。"""
        if isinstance(start, int) and isinstance(end, int):
            if 0 <= start < end <= len(text) and text[start:end] == value:
                return True, start, end
        idx = text.find(value)
        if idx >= 0:
            return True, idx, idx + len(value)
        # 宽松：去掉空白后再找（容忍模型混入空格）
        compact = re.sub(r"\s+", "", value)
        if compact:
            idx = text.find(compact)
            if idx >= 0:
                return True, idx, idx + len(value)
        return False, 0, 0

    # ---------------- 文件：PDF / 图片 ----------------

    @staticmethod
    def extract_pdf_text(data: bytes) -> dict:
        """提取 PDF 文本；若是无文字层的扫描件/图片型 PDF（提取为空），则逐页渲染 OCR。

        返回结构含 scanned 标记：带文字层 scanned=False；扫描件 scanned=True（文本来自 OCR）。
        """
        try:
            import fitz  # pymupdf
        except ImportError as e:  # pragma: no cover
            raise RuntimeError("未安装 pymupdf，无法处理 PDF（pip install pymupdf）") from e
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        full = []
        for p in doc:
            t = p.get_text()
            pages.append(t)
            full.append(t)
        doc.close()
        text = "\n".join(full).strip()
        if text:
            return {"pages": pages, "text": "\n".join(full),
                    "pageCount": len(pages), "kind": "pdf", "scanned": False}
        # 无文字层 → 扫描件：逐页渲染 OCR
        return DesensitizeService._ocr_pdf(data)

    @staticmethod
    def _ocr_pdf(data: bytes) -> dict:
        """扫描件 PDF：每页渲染成图后 OCR，返回每页文本与全文（scanned=True）。"""
        import fitz
        if not _ensure_tesseract():
            raise RuntimeError("未检测到 Tesseract OCR，无法识别扫描件 PDF（请安装 Tesseract-OCR 并含中文语言包）")
        import pytesseract
        from PIL import Image
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        full = []
        for p in doc:
            pix = p.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            try:
                t = pytesseract.image_to_string(img, lang=_OCR_LANG)
            except Exception:
                t = ""
            pages.append(t)
            full.append(t)
        doc.close()
        return {"pages": pages, "text": "\n".join(full), "pageCount": len(pages),
                "kind": "pdf", "scanned": True}

    @staticmethod
    def extract_doc_text(data: bytes, ext: str) -> dict:
        """提取办公文档文本：.docx（python-docx 完整解析）/ .doc（尽力）/ .txt / .md 直接读取。

        返回 {"text", "kind"}；无法提取（如 .doc 二进制无清晰文本）抛 RuntimeError。"""
        ext = (ext or "").lower()
        if ext == "docx":
            try:
                import docx
            except ImportError as e:  # pragma: no cover
                raise RuntimeError("未安装 python-docx，无法解析 .docx（pip install python-docx）") from e
            d = docx.Document(io.BytesIO(data))
            paras = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    paras.append("\t".join(c.text for c in row.cells))
            return {"text": "\n".join(paras), "kind": "docx"}
        if ext == "doc":
            t = _doc_binary_text(data)
            if not t.strip():
                raise RuntimeError("无法从 .doc 提取文本（旧二进制格式），请另存为 .docx 或 .txt 后上传")
            return {"text": t, "kind": "doc"}
        if ext in ("txt", "md", "markdown", "text"):
            return {"text": data.decode("utf-8", errors="replace"), "kind": "text"}
        raise RuntimeError(f"不支持的文件类型: {ext or '(无扩展名)'}")

    def extract_image_text(self, data: bytes) -> dict:
        """图片 OCR（pytesseract + 本机 Tesseract）：返回文本与词级坐标表；OCR 不可用返回空。"""
        if not _ensure_tesseract():
            return {"text": "", "words": [], "ocr": False, "kind": "image"}
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return {"text": "", "words": [], "ocr": False, "kind": "image"}
        try:
            img = Image.open(io.BytesIO(data))
            data_boxes = pytesseract.image_to_data(img, lang=_OCR_LANG,
                                                   output_type=pytesseract.Output.DICT)
            txt = pytesseract.image_to_string(img, lang=_OCR_LANG)
        except Exception:
            return {"text": "", "words": [], "ocr": False, "kind": "image"}
        words = []
        n = len(data_boxes.get("text", []))
        for i in range(n):
            w = (data_boxes.get("text") or [""])[i]
            if not w or not w.strip():
                continue
            words.append({"word": w.strip(),
                          "x": data_boxes["left"][i], "y": data_boxes["top"][i],
                          "w": data_boxes["width"][i], "h": data_boxes["height"][i]})
        return {"text": txt, "words": words, "ocr": True, "kind": "image"}

    # ---------------- 工具：PDF / 图片涂黑 ----------------

    @staticmethod
    def redact_pdf(data: bytes, values: list[str]) -> bytes:
        """把 PDF 中每个敏感串出现的位置涂成黑色块，返回新 PDF。

        带文字层 PDF 按文字定位涂黑；无文字层的扫描件 PDF 走 OCR 词坐标定位，
        在渲染图上画黑块并重建 PDF（_redact_scanned_pdf）。
        """
        try:
            import fitz
        except ImportError as e:  # pragma: no cover
            raise RuntimeError("未安装 pymupdf，无法涂黑 PDF（pip install pymupdf）") from e
        doc = fitz.open(stream=data, filetype="pdf")
        has_text = any(p.get_text().strip() for p in doc)
        doc.close()
        if not has_text:
            return DesensitizeService._redact_scanned_pdf(data, values)
        doc = fitz.open(stream=data, filetype="pdf")
        for page in doc:
            rects = []
            pr = page.rect
            for v in values:
                if not v:
                    continue
                for r in page.search_for(v):
                    clip = fitz.Rect(r) & pr
                    if not clip.is_empty:
                        rects.append(clip)
            if rects:
                for clip in rects:
                    page.add_redact_annot(clip, fill=(0, 0, 0))
                page.apply_redactions()
        buf = io.BytesIO()
        doc.save(buf, garbage=3, deflate=True)
        doc.close()
        return buf.getvalue()

    @staticmethod
    def _redact_scanned_pdf(data: bytes, values: list[str]) -> bytes:
        """扫描件 PDF 涂黑：每页渲染成图 → OCR 词坐标 → 对敏感串占用矩形画黑块 → 重建 PDF。"""
        import fitz
        if not _ensure_tesseract():
            raise RuntimeError("未检测到 Tesseract OCR，无法涂黑扫描件 PDF（请安装 Tesseract-OCR 并含中文语言包）")
        import pytesseract
        from PIL import Image, ImageDraw
        doc = fitz.open(stream=data, filetype="pdf")
        out = fitz.open()
        matrix = fitz.Matrix(2, 2)
        for page in doc:
            pix = page.get_pixmap(matrix=matrix)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            draw = ImageDraw.Draw(img)
            try:
                boxes = pytesseract.image_to_data(img, lang=_OCR_LANG,
                                                  output_type=pytesseract.Output.DICT)
            except Exception:
                boxes = {"text": [], "left": [], "top": [], "width": [], "height": []}
            words = []
            n = len(boxes.get("text", []) or [])
            for i in range(n):
                w = (boxes.get("text") or [""])[i]
                if not w or not str(w).strip():
                    continue
                words.append([str(w).strip(),
                              int(boxes["left"][i]), int(boxes["top"][i]),
                              int(boxes["width"][i]), int(boxes["height"][i])])
            rects = DesensitizeService._find_value_rects(words, values)
            for (x0, y0, x1, y1) in rects:
                draw.rectangle([x0, y0, x1, y1], fill=(0, 0, 0))
            npage = out.new_page(width=page.rect.width, height=page.rect.height)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            npage.insert_image(page.rect, stream=buf.getvalue())
        doc.close()
        outbuf = io.BytesIO()
        out.save(outbuf, garbage=3, deflate=True)
        out.close()
        return outbuf.getvalue()

    @staticmethod
    def _find_value_rects(words: list[list], values: list[str]) -> list[tuple]:
        """在 OCR word 序列（[[text,x,y,w,h],...]）中定位敏感串出现的矩形并集（渲染图像素坐标）。

        对每个 value 在「去空格的 word 拼接文本」里找子串，取覆盖该子串的所有 word 的
        外接矩形；同 value 多处出现都返回。
        """
        rects_out = []
        if not words:
            return rects_out
        texts = [w[0] for w in words]
        joined = "".join(texts)
        span = []
        pos = 0
        for t in texts:
            span.append((pos, pos + len(t)))
            pos += len(t)
        for value in values:
            if not value:
                continue
            compact = re.sub(r"\s+", "", value)
            if not compact:
                continue
            idx = joined.find(compact)
            while idx >= 0:
                endi = idx + len(compact)
                covers = [i for i, (s, e) in enumerate(span) if s < endi and e > idx]
                if covers:
                    xs = [words[i][1] for i in covers]
                    ys = [words[i][2] for i in covers]
                    x1s = [words[i][1] + words[i][3] for i in covers]
                    y1s = [words[i][2] + words[i][4] for i in covers]
                    rects_out.append((min(xs), min(ys), max(x1s), max(y1s)))
                idx = joined.find(compact, idx + 1)
        return rects_out

    @staticmethod
    def redact_image(data: bytes, regions: Optional[list[list[int]]] = None,
                     full: bool = False) -> bytes:
        """把图片指定矩形或整张涂成黑块，返回新图字节（保留原格式；PNG 兜底）。"""
        try:
            from PIL import Image, ImageDraw
        except ImportError as e:  # pragma: no cover
            raise RuntimeError("未安装 Pillow，无法涂黑图片（pip install Pillow）") from e
        img = Image.open(io.BytesIO(data))
        fmt = (img.format or "PNG").upper()
        if fmt in ("", "JPEG", "JPG"):
            rgb = img.convert("RGB")
        else:
            rgb = img.convert("RGBA") if img.mode in ("RGBA", "LA") or "A" in (img.getbands() or ()) else img.convert("RGB")
        draw = ImageDraw.Draw(rgb)
        if full or not regions:
            draw.rectangle([0, 0, rgb.width, rgb.height], fill=(0, 0, 0))
        else:
            for reg in regions:
                if len(reg) != 4:
                    continue
                x0, y0, x1, y1 = reg
                draw.rectangle([x0, y0, x1, y1], fill=(0, 0, 0))
        buf = io.BytesIO()
        rgb.save(buf, format="PNG" if fmt in ("JPEG", "JPG") else fmt)
        return buf.getvalue()
